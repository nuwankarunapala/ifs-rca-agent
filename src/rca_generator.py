"""
rca_generator.py — Generates a formatted Word (.docx) RCA report.

Document structure (9 sections matching Claude's output):
  Cover Page
  1. Executive Summary
  2. Scope & Impact
  3. Timeline of Events
  4. Technical Analysis
  5. Root Cause
  6. Corrective Actions
  7. Preventive Actions
  8. Validation Plan
  9. Appendix  (top-5 most diagnostic log excerpts)
"""

import re
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from src.log_parser import LogError

# ---------------------------------------------------------------------------
# Heading detection — matches Claude's Markdown output
# ---------------------------------------------------------------------------
_H1_RE  = re.compile(r"^##\s+(.+)")          # ## 1. Title
_H2_RE  = re.compile(r"^###\s+(.+)")         # ### 4a. Sub-title
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def generate_rca_document(
    errors: List[LogError],
    user_context: Dict[str, str],
    analysis: str,
) -> str:
    """
    Generate the Word RCA report.

    Args:
        errors:       Parsed LogError objects.
        user_context: Incident context from the user.
        analysis:     Claude's 9-section Markdown RCA text.

    Returns:
        Absolute path of the saved .docx file (string).
    """
    output_dir = Path("output")
    output_dir.mkdir(parents=True, exist_ok=True)

    now = datetime.now(timezone.utc)
    filename = f"RCA_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = output_dir / filename

    doc = Document()
    _set_default_font(doc)

    # ------------------------------------------------------------------
    # Cover page
    # ------------------------------------------------------------------
    _cover_page(doc, user_context, now, len(errors))

    # ------------------------------------------------------------------
    # Error Inventory table (inside Section 3 / before main analysis)
    # ------------------------------------------------------------------
    doc.add_page_break()
    _error_inventory(doc, errors)

    # ------------------------------------------------------------------
    # Main analysis — parsed from Claude's Markdown output
    # ------------------------------------------------------------------
    doc.add_page_break()
    _render_analysis(doc, analysis)

    doc.save(output_path)
    return str(output_path.resolve())


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _cover_page(doc: Document, ctx: Dict[str, str], now: datetime, total_errors: int) -> None:
    title = doc.add_heading("Root Cause Analysis Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    meta = [
        ("Generated",         now.strftime("%Y-%m-%d %H:%M:%S UTC")),
        ("Environment",       ctx.get("environment", "N/A")),
        ("Incident Time",     ctx.get("incident_time", "N/A")),
        ("Affected Services", ctx.get("affected_services", "N/A")),
        ("Recent Changes",    ctx.get("recent_changes", "None")),
        ("Total Log Errors",  str(total_errors)),
    ]

    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(meta):
        cells = tbl.rows[i].cells
        cells[0].text = label
        _bold_run(cells[0])
        cells[1].text = value


# ---------------------------------------------------------------------------
# Error inventory table
# ---------------------------------------------------------------------------

def _error_inventory(doc: Document, errors: List[LogError]) -> None:
    doc.add_heading("Error Inventory", level=1)

    # Summary by type
    by_type = Counter(e.error_type for e in errors)
    summary = doc.add_table(rows=1, cols=2)
    summary.style = "Table Grid"
    _header_row(summary.rows[0].cells, ["Error Type", "Count"])
    for etype, count in by_type.most_common():
        row = summary.add_row().cells
        row[0].text = etype
        row[1].text = str(count)

    doc.add_paragraph()

    # Detail table (max 50 rows)
    doc.add_paragraph().add_run("Detailed Error Log (first 50 entries)").bold = True
    detail = doc.add_table(rows=1, cols=5)
    detail.style = "Table Grid"
    _header_row(
        detail.rows[0].cells,
        ["Timestamp", "Error Type", "Pod", "Source Type", "Source File"],
    )
    for err in errors[:50]:
        row = detail.add_row().cells
        row[0].text = err.timestamp
        row[1].text = err.error_type
        row[2].text = err.pod_name
        row[3].text = err.file_type
        row[4].text = err.source_file


# ---------------------------------------------------------------------------
# Main analysis renderer
# ---------------------------------------------------------------------------

def _render_analysis(doc: Document, analysis: str) -> None:
    """
    Walk Claude's Markdown output line by line and render into the Word doc.
    Handles: ## headings, ### sub-headings, | tables, ``` code blocks,
    bullet lists, and plain paragraphs with **bold** inline.
    """
    in_code   = False
    code_buf: List[str] = []
    in_table  = False
    tbl_obj   = None

    lines = analysis.splitlines()

    for line in lines:
        # ---- code block ----
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                _add_code_block(doc, "\n".join(code_buf))
                code_buf = []
            continue
        if in_code:
            code_buf.append(line)
            continue

        # ---- table rows ----
        if line.strip().startswith("|"):
            cells_raw = [c.strip() for c in line.strip().strip("|").split("|")]
            if all(re.match(r"^[-: ]+$", c) for c in cells_raw):
                # separator row — skip
                continue
            if not in_table:
                in_table = True
                tbl_obj = doc.add_table(rows=1, cols=len(cells_raw))
                tbl_obj.style = "Table Grid"
                _header_row(tbl_obj.rows[0].cells, cells_raw)
            else:
                row = tbl_obj.add_row().cells
                for i, text in enumerate(cells_raw):
                    if i < len(row):
                        row[i].text = text
            continue
        else:
            in_table = False
            tbl_obj = None

        # ---- headings ----
        m = _H1_RE.match(line)
        if m:
            doc.add_heading(m.group(1), level=1)
            continue

        m = _H2_RE.match(line)
        if m:
            doc.add_heading(m.group(1), level=2)
            continue

        # ---- bullet lists ----
        stripped = line.lstrip()
        if stripped.startswith(("- ", "* ", "+ ")):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_bold(p, text)
            continue

        # Numbered list  "1. " / "2. "
        nm = re.match(r"^\d+\.\s+(.*)", stripped)
        if nm:
            p = doc.add_paragraph(style="List Number")
            _add_inline_bold(p, nm.group(1))
            continue

        # ---- blank line ----
        if not line.strip():
            doc.add_paragraph()
            continue

        # ---- plain paragraph ----
        p = doc.add_paragraph()
        _add_inline_bold(p, line)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_code_block(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


def _add_inline_bold(para, text: str) -> None:
    """Render **bold** markers inside a paragraph."""
    parts = _BOLD_RE.split(text)
    for i, part in enumerate(parts):
        run = para.add_run(part)
        run.bold = (i % 2 == 1)


def _header_row(cells, labels: List[str]) -> None:
    for cell, label in zip(cells, labels):
        cell.text = label
        _bold_run(cell)


def _bold_run(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def generate_health_report(
    errors: List[LogError],
    analysis: str,
    logs_dir: str = "",
) -> str:
    """
    Generate a Word health report from Claude's 7-section health assessment.

    Args:
        errors:    All LogError objects detected in the logs.
        analysis:  Claude's 7-section Markdown health report text.
        logs_dir:  Path to the logs directory (shown on cover page).

    Returns:
        Absolute path of the saved .docx file (string).
    """
    output_dir = Path("output")
    output_dir.mkdir(parents=True, exist_ok=True)

    now = datetime.now(timezone.utc)
    filename = f"HealthReport_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = output_dir / filename

    doc = Document()
    _set_default_font(doc)

    _health_cover_page(doc, now, errors, logs_dir)

    doc.add_page_break()
    _health_event_summary(doc, errors)

    doc.add_page_break()
    _render_analysis(doc, analysis)

    doc.save(output_path)
    return str(output_path.resolve())


def _health_cover_page(
    doc: Document,
    now: datetime,
    errors: List[LogError],
    logs_dir: str,
) -> None:
    title = doc.add_heading("Cluster Health Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("IFS / Kubernetes — Proactive Health Assessment")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    affected_pods = list({e.pod_name for e in errors if e.pod_name != "unknown"})
    pod_list = ", ".join(sorted(affected_pods)[:10]) or "None detected"
    if len(affected_pods) > 10:
        pod_list += f" (+{len(affected_pods) - 10} more)"

    by_type = Counter(e.error_type for e in errors)
    top_issues = ", ".join(f"{t} ({c})" for t, c in by_type.most_common(5)) or "None"

    meta = [
        ("Report Generated",  now.strftime("%Y-%m-%d %H:%M:%S UTC")),
        ("Analysis Type",     "Proactive Health Assessment"),
        ("Logs Source",       logs_dir or "logs/"),
        ("Total Events Found", str(len(errors))),
        ("Unique Event Types", str(len(by_type))),
        ("Top Issues",        top_issues),
        ("Pods With Events",  pod_list),
    ]

    tbl = doc.add_table(rows=len(meta), cols=2)
    tbl.style = "Table Grid"
    for i, (label, value) in enumerate(meta):
        cells = tbl.rows[i].cells
        cells[0].text = label
        _bold_run(cells[0])
        cells[1].text = value


def _health_event_summary(doc: Document, errors: List[LogError]) -> None:
    doc.add_heading("Event Summary", level=1)

    by_type = Counter(e.error_type for e in errors)
    summary = doc.add_table(rows=1, cols=3)
    summary.style = "Table Grid"
    _header_row(summary.rows[0].cells, ["Event Type", "Count", "Severity"])

    _sev_map: Dict[str, str] = {}
    for e in errors:
        if e.error_type not in _sev_map or e.severity < _sev_map[e.error_type]:
            _sev_map[e.error_type] = e.severity

    for etype, count in by_type.most_common():
        row = summary.add_row().cells
        row[0].text = etype
        row[1].text = str(count)
        row[2].text = _sev_map.get(etype, "")

    doc.add_paragraph()

    doc.add_paragraph().add_run("Detailed Event Log (first 50 entries)").bold = True
    detail = doc.add_table(rows=1, cols=5)
    detail.style = "Table Grid"
    _header_row(
        detail.rows[0].cells,
        ["Timestamp", "Event Type", "Pod", "Source Type", "Source File"],
    )
    for err in errors[:50]:
        row = detail.add_row().cells
        row[0].text = err.timestamp
        row[1].text = err.error_type
        row[2].text = err.pod_name
        row[3].text = err.file_type
        row[4].text = err.source_file


def _set_default_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    font  = style.font
    font.name = "Calibri"
    font.size = Pt(11)
